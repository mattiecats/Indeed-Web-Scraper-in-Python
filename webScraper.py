#!/usr/bin/env python3
import docx
import requests
from bs4 import BeautifulSoup
from docx.enum.section import WD_SECTION_START


def htmlRequest(input): #function that parses URL for input from Indeed Listing
    r = requests.get(input) 
    soup = BeautifulSoup(r.text, features='html.parser')
    print(soup.title.string)
    for match in soup.find_all('div', id='jobDescriptionText'):
        print(match.get_text())
    doc = docx.Document()
    heading = doc.add_heading(soup.title.string, 0)
    paragraph = doc.add_paragraph(match.get_text())
    new_section = doc.add_section(WD_SECTION_START.ODD_PAGE)
    doc.save('Indeed.docx')


def main(): #main function to loop to write to docx file for output
    while input != '-1':
        x = input('Enter the URL here or 0 to quit: ')
        if x == 0:
            break
        else:
            htmlRequest(x)


if __name__ == '__main__': main()
